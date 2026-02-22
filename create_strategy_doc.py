#!/usr/bin/env python3
"""
生成堕落天使策略Word文档
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_table(doc, headers, rows, header_color='4472C4'):
    """添加表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # 表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], header_color)
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    # 数据行
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)

    doc.add_paragraph()
    return table

def create_strategy_doc():
    doc = Document()

    # 设置中文字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # ========== 封面 ==========
    title = doc.add_heading('堕落天使策略', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Fallen Angel Strategy')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    core = doc.add_paragraph()
    core.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = core.add_run('核心逻辑：热门板块 + 前期暴涨 + 当前超跌 + 非成分股 = 反弹爆发力最强')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(192, 0, 0)

    doc.add_paragraph()
    date_p = doc.add_paragraph(f'更新日期: {datetime.date.today().strftime("%Y-%m-%d")}')
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ========== 目录 ==========
    add_heading(doc, '目录', 1)
    toc = [
        '一、策略定义',
        '二、选股条件',
        '三、目标股票池（AI核心板块）',
        '四、扩展板块（7大热门板块）',
        '五、买入时机',
        '六、持仓与卖出',
        '七、仓位管理',
        '八、实战案例',
        '九、每日操作清单',
        '十、风险提示',
    ]
    for item in toc:
        doc.add_paragraph(item, style='List Number')

    doc.add_page_break()

    # ========== 一、策略定义 ==========
    add_heading(doc, '一、策略定义', 1)

    p = doc.add_paragraph()
    run = p.add_run('一句话：')
    run.bold = True
    p.add_run('买入热门板块中跌惨了的前期强势股（非恒科成分股），等利好消息爆拉。')

    add_heading(doc, '策略标签', 2)
    add_table(doc, ['标签', '说明'], [
        ['类型', '超跌反弹 + 事件驱动'],
        ['周期', '波段（1-10天）'],
        ['频率', '低频（每周1-3次）'],
        ['风险', '中高'],
        ['收益目标', '单笔15-30%'],
    ])

    add_heading(doc, '大师背书', 2)
    add_table(doc, ['大师', '相关理论'], [
        ['William O\'Neil', 'CANSLIM - 调整时买领涨股回调'],
        ['Mark Minervini', 'VCP - 波动收缩后买前期强势股'],
        ['Jesse Livermore', '关键位回调，在支撑位买入'],
    ])

    # ========== 二、选股条件 ==========
    add_heading(doc, '二、选股条件（全部满足）', 1)

    add_heading(doc, '必须条件 ✅', 2)
    add_table(doc, ['#', '条件', '具体标准', '原因'], [
        ['1', '热门板块', 'AI大模型/机器人/芯片/生物科技', '资金持续关注'],
        ['2', '前期暴涨', '过去60天有单日涨幅>10%，或30天累计涨>30%', '证明市场认可'],
        ['3', '当前超跌', '从近期高点回撤>25%，或连续下跌>3天', '筹码松动，抛压出清'],
        ['4', '非恒科成分股', '不在恒生科技指数30只成分股内', '保证弹性'],
        ['5', '流动性', '日成交额>3000万港元', '能买能卖'],
    ], '2E7D32')

    add_heading(doc, '排除条件 ❌', 2)
    add_table(doc, ['条件', '原因'], [
        ['恒科指数成分股', '市值大、弹性差'],
        ['有明确利空消息', '可能继续跌'],
        ['财务造假/退市风险', '血本无归'],
        ['解禁期刚过', '抛压未出清'],
        ['大股东减持中', '信心不足'],
    ], 'C00000')

    add_heading(doc, '恒科成分股名单（排除）', 2)
    p = doc.add_paragraph()
    p.add_run('腾讯00700、阿里09988、美团03690、京东09618、网易09999、百度09888、小米01810、快手01024、联想00992、中芯国际00981、华虹半导体01347、比亚迪电子00285、舜宇光学02382、瑞声科技02018、金蝶国际00268、微盟集团02013、明源云00909、商汤00020、阅文集团00772、哔哩哔哩09626、携程09961、贝壳02423、蔚来09866、理想02015、小鹏09868、金山软件03888、万国数据09698、金山云03896、微博09898、知乎02390').font.size = Pt(9)

    # ========== 三、目标股票池 ==========
    add_heading(doc, '三、目标股票池（AI核心板块）', 1)

    add_heading(doc, '🔥 AI大模型概念（弹性最高）', 2)
    add_table(doc, ['股票', '代码', '简介', '弹性'], [
        ['MiniMax', '00100.HK', '全球最快IPO的AI公司，认购超1800倍', '⭐⭐⭐⭐⭐'],
        ['智谱AI', '02513.HK', '大模型第一股，清华系，市值500亿+', '⭐⭐⭐⭐⭐'],
        ['云知声', '09678.HK', '港股AGI第一股，AI语音龙头', '⭐⭐⭐⭐'],
    ], 'FF6B00')

    add_heading(doc, '💾 AI芯片/GPU概念', 2)
    add_table(doc, ['股票', '代码', '简介', '弹性'], [
        ['壁仞科技', '06082.HK', '港股GPU第一股，首日涨76%', '⭐⭐⭐⭐⭐'],
        ['天数智芯', '09903.HK', '国产通用GPU，聚焦训推', '⭐⭐⭐⭐⭐'],
        ['地平线机器人', '09660.HK', '智驾芯片+机器人芯片', '⭐⭐⭐⭐'],
    ], '7030A0')

    add_heading(doc, '🤖 机器人概念', 2)
    add_table(doc, ['股票', '代码', '类型', '弹性'], [
        ['优必选', '09880.HK', '人形机器人第一股，专利全球第一', '⭐⭐⭐⭐'],
        ['越疆', '02432.HK', '协作机器人第一股，出货量中国第一', '⭐⭐⭐⭐'],
        ['微创机器人', '02252.HK', '手术机器人龙头，图迈系统', '⭐⭐⭐⭐'],
        ['精锋医疗', '02675.HK', '腔镜手术机器人，首日涨36%', '⭐⭐⭐⭐'],
        ['极智嘉', '02590.HK', '全球仓储机器人龙头，蚂蚁系', '⭐⭐⭐'],
    ], '0070C0')

    add_heading(doc, '💊 AI制药/生物科技', 2)
    add_table(doc, ['股票', '代码', '简介', '弹性'], [
        ['英矽智能', '02658.HK', 'AI制药第一股，端到端能力', '⭐⭐⭐⭐'],
        ['晶泰科技', '02259.HK', 'AI+CRO，量子物理药物晶型', '⭐⭐⭐⭐'],
        ['讯飞医疗', '02506.HK', 'AI赋能医疗解决方案领导者', '⭐⭐⭐'],
        ['医渡科技', '02158.HK', '医疗AI平台，字节生态合作', '⭐⭐⭐'],
    ], '00B050')

    add_heading(doc, '🧠 AI平台', 2)
    add_table(doc, ['股票', '代码', '简介', '弹性'], [
        ['第四范式', '06682.HK', '企业级AI平台，决策式AI', '⭐⭐⭐⭐'],
    ], '4472C4')

    doc.add_page_break()

    # ========== 四、扩展板块 ==========
    add_heading(doc, '四、扩展板块（7大热门板块）', 1)

    add_heading(doc, '🚁 低空经济/eVTOL', 2)
    p = doc.add_paragraph()
    run = p.add_run('催化剂：')
    run.bold = True
    p.add_run('六城eVTOL试点（深圳/杭州/合肥/苏州/成都/重庆），2026年万亿市场')

    add_table(doc, ['股票', '代码', '简介', '弹性'], [
        ['小鹏汽车-W', '09868.HK', '小鹏汇天"陆地航母"，2026量产', '⭐⭐⭐⭐'],
        ['吉利汽车', '00175.HK', '沃飞长空AE200，签约100架', '⭐⭐⭐'],
        ['广汽集团', '02238.HK', 'GOVE飞行汽车已首飞', '⭐⭐⭐'],
        ['亚太卫星', '01045.HK', '低空通信基础设施', '⭐⭐⭐'],
    ], '00B0F0')

    add_heading(doc, '⚡ 新能源/固态电池', 2)
    p = doc.add_paragraph()
    run = p.add_run('催化剂：')
    run.bold = True
    p.add_run('工信部全固态电池2027年攻关目标')

    add_table(doc, ['股票', '代码', '简介', '弹性'], [
        ['赣锋锂业', '01772.HK', '全球唯一硫化物电解质量产', '⭐⭐⭐⭐'],
        ['天齐锂业', '09696.HK', '锂资源龙头', '⭐⭐⭐⭐'],
        ['信义光能', '00968.HK', '光伏玻璃龙头', '⭐⭐⭐'],
        ['信义能源', '03868.HK', '光伏电站运营', '⭐⭐⭐'],
    ], '00B050')

    add_heading(doc, '🥇 黄金/有色金属', 2)
    p = doc.add_paragraph()
    run = p.add_run('催化剂：')
    run.bold = True
    p.add_run('黄金目标$5000/盎司，稀土出口管制')

    add_table(doc, ['股票', '代码', '简介', '弹性'], [
        ['紫金矿业', '02899.HK', '综合矿业龙头，黄金+铜', '⭐⭐⭐⭐'],
        ['洛阳钼业', '03993.HK', '钴钼资源全球前列', '⭐⭐⭐⭐'],
        ['江西铜业', '00358.HK', '铜业龙头', '⭐⭐⭐'],
        ['中国黄金国际', '02099.HK', '央企黄金矿业', '⭐⭐⭐'],
        ['招金矿业', '01818.HK', '黄金矿业', '⭐⭐⭐'],
    ], 'FFD700')

    add_heading(doc, '🎮 消费出海/潮玩', 2)
    p = doc.add_paragraph()
    run = p.add_run('催化剂：')
    run.bold = True
    p.add_run('海外IP授权+东南亚/欧美扩张')

    add_table(doc, ['股票', '代码', '简介', '弹性'], [
        ['泡泡玛特', '09992.HK', '潮玩第一股，Labubu全球爆火', '⭐⭐⭐⭐⭐'],
        ['名创优品', '09896.HK', '出海零售龙头', '⭐⭐⭐⭐'],
        ['老铺黄金', '06181.HK', '国潮黄金饰品', '⭐⭐⭐⭐'],
        ['安踏体育', '02020.HK', '运动品牌出海', '⭐⭐⭐'],
    ], 'FF69B4')

    add_heading(doc, '☢️ 核电/铀矿', 2)
    p = doc.add_paragraph()
    run = p.add_run('催化剂：')
    run.bold = True
    p.add_run('两大铀供应商下调2026产量10%')

    add_table(doc, ['股票', '代码', '简介', '弹性'], [
        ['中广核矿业', '01164.HK', '铀资源龙头，近期涨超9%', '⭐⭐⭐⭐'],
        ['中广核电力', '01816.HK', '核电运营龙头，28台在运', '⭐⭐⭐'],
        ['中广核新能源', '01811.HK', '清洁能源平台', '⭐⭐⭐'],
    ], '9400D3')

    add_heading(doc, '🚀 军工/商业航天', 2)
    p = doc.add_paragraph()
    run = p.add_run('催化剂：')
    run.bold = True
    p.add_run('SpaceX+xAI合并，太空算力时代')

    add_table(doc, ['股票', '代码', '简介', '弹性'], [
        ['中航科工', '02357.HK', '航空装备核心平台', '⭐⭐⭐⭐'],
        ['亚太卫星', '01045.HK', '卫星通信运营', '⭐⭐⭐'],
    ], '808080')

    add_heading(doc, '⚡ 电网/特高压', 2)
    p = doc.add_paragraph()
    run = p.add_run('催化剂：')
    run.bold = True
    p.add_run('国家电网"十五五"投资4万亿，增长40%')

    add_table(doc, ['股票', '代码', '简介', '弹性'], [
        ['东方电气', '01072.HK', '发电设备龙头', '⭐⭐⭐'],
        ['上海电气', '02727.HK', '综合电气装备', '⭐⭐⭐'],
        ['哈尔滨电气', '01133.HK', '发电设备三巨头', '⭐⭐⭐'],
        ['长飞光纤', '06869.HK', '光纤光缆龙头', '⭐⭐⭐'],
        ['威胜控股', '03393.HK', '智能电表配电', '⭐⭐⭐'],
    ], '4472C4')

    doc.add_page_break()

    # ========== 股票池速查表 ==========
    add_heading(doc, '📋 股票池速查表（共42只）', 1)

    p = doc.add_paragraph()
    run = p.add_run('【高弹性板块 - 堕落天使首选】')
    run.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)

    doc.add_paragraph('AI大模型: 00100(MiniMax) | 02513(智谱) | 09678(云知声)')
    doc.add_paragraph('AI芯片: 06082(壁仞) | 09903(天数智芯) | 09660(地平线)')
    doc.add_paragraph('机器人: 09880(优必选) | 02432(越疆) | 02252(微创) | 02675(精锋) | 02590(极智嘉)')
    doc.add_paragraph('AI制药: 02658(英矽智能) | 02259(晶泰) | 02506(讯飞医疗) | 02158(医渡)')
    doc.add_paragraph('AI平台: 06682(第四范式)')
    doc.add_paragraph('潮玩出海: 09992(泡泡玛特) | 09896(名创) | 06181(老铺黄金)')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('【中弹性板块 - 波段机会】')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 112, 192)

    doc.add_paragraph('低空经济: 09868(小鹏) | 00175(吉利) | 02238(广汽) | 01045(亚太卫星)')
    doc.add_paragraph('新能源: 01772(赣锋) | 09696(天齐) | 00968(信义光能)')
    doc.add_paragraph('有色金属: 02899(紫金) | 03993(洛钼) | 00358(江西铜业)')
    doc.add_paragraph('核电铀矿: 01164(中广核矿业) | 01816(中广核电力)')
    doc.add_paragraph('军工航天: 02357(中航科工)')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('【低弹性板块 - 稳健配置】')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 176, 80)

    doc.add_paragraph('电网设备: 01072(东方电气) | 02727(上海电气) | 01133(哈电) | 06869(长飞)')

    doc.add_page_break()

    # ========== 五、买入时机 ==========
    add_heading(doc, '五、买入时机', 1)

    add_heading(doc, '最佳买入信号', 2)
    add_table(doc, ['信号', '权重', '说明'], [
        ['地量企稳', '⭐⭐⭐⭐⭐', '成交量萎缩到20日均量的50%以下'],
        ['长下影线', '⭐⭐⭐⭐', '说明下方有承接'],
        ['跌到整数关口', '⭐⭐⭐⭐', '如10元、20元、50元'],
        ['跌到前低支撑', '⭐⭐⭐⭐', '技术面支撑'],
        ['板块利好消息', '⭐⭐⭐⭐⭐', '催化剂来了'],
        ['大盘企稳反弹', '⭐⭐⭐', '系统性风险解除'],
    ])

    add_heading(doc, '买入时间窗口', 2)
    add_table(doc, ['时间', '操作'], [
        ['09:30-10:00', '观察，不急着买'],
        ['10:00-11:00 ⭐', '确认企稳后首次建仓（50%仓位）'],
        ['14:00-15:00', '若继续企稳，加仓（50%仓位）'],
        ['15:30-16:00 ⭐', '若有利好传闻，可博隔夜'],
    ])

    add_heading(doc, '分批建仓策略', 2)
    p = doc.add_paragraph()
    run = p.add_run('第一笔：50%仓位，在确认企稳时买入\n第二笔：50%仓位，在确认反弹或利好出现时加仓\n\n')
    run = p.add_run('⚠️ 不要一把梭！')
    run.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)

    # ========== 六、持仓与卖出 ==========
    add_heading(doc, '六、持仓与卖出', 1)

    add_heading(doc, '持仓周期', 2)
    add_table(doc, ['情况', '持仓时间'], [
        ['有利好催化', '1-3天（快进快出）'],
        ['无明确催化', '3-7天（等待）'],
        ['超过7天无反弹', '考虑止损'],
    ])

    add_heading(doc, '止盈规则 📈', 2)
    add_table(doc, ['条件', '操作'], [
        ['反弹10%', '止盈30%仓位'],
        ['反弹15%', '止盈50%仓位'],
        ['反弹20%+', '移动止盈，追踪'],
        ['利好兑现', '全部止盈'],
        ['放量冲高回落', '立即止盈'],
    ], '2E7D32')

    add_heading(doc, '止损规则 📉', 2)
    add_table(doc, ['条件', '操作'], [
        ['跌破买入价8%', '止损'],
        ['出现利空消息', '立即止损'],
        ['板块整体走弱', '减仓50%'],
        ['持仓超7天无反弹', '止损'],
        ['大盘系统性下跌', '减仓'],
    ], 'C00000')

    # ========== 七、仓位管理 ==========
    add_heading(doc, '七、仓位管理', 1)

    add_heading(doc, '资金分配（5万本金）', 2)
    add_table(doc, ['项目', '比例', '金额'], [
        ['单只股票最大仓位', '30%', '15,000元'],
        ['同板块最多', '2只', '30,000元'],
        ['总仓位上限', '60%', '30,000元'],
        ['现金保留', '40%', '20,000元'],
    ])

    add_heading(doc, '风控红线 🚨', 2)
    p = doc.add_paragraph()
    run = p.add_run('单笔最大亏损：')
    run.bold = True
    p.add_run('8% × 15000 = 1200元\n')
    run = p.add_run('单日最大亏损：')
    run.bold = True
    p.add_run('2000元 → 达到就停止\n')
    run = p.add_run('连续亏损3次：')
    run.bold = True
    p.add_run('暂停3天，复盘策略')

    doc.add_page_break()

    # ========== 八、实战案例 ==========
    add_heading(doc, '八、实战案例', 1)

    add_heading(doc, '案例1：MiniMax (00100)', 2)

    p = doc.add_paragraph()
    run = p.add_run('背景：')
    run.bold = True
    doc.add_paragraph('• AI大模型龙头，前期暴涨200%+')
    doc.add_paragraph('• 回调40%，市场恐慌')
    doc.add_paragraph('• 非恒科成分股，市值小')

    p = doc.add_paragraph()
    run = p.add_run('买入：')
    run.bold = True
    doc.add_paragraph('• 买入价：30元（回调后企稳）')
    doc.add_paragraph('• 仓位：15000元')
    doc.add_paragraph('• 触发条件：地量+长下影线')

    p = doc.add_paragraph()
    run = p.add_run('催化剂：')
    run.bold = True
    doc.add_paragraph('• 第二天传出与字节合作消息')

    p = doc.add_paragraph()
    run = p.add_run('结果：')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 176, 80)
    doc.add_paragraph('• 卖出价：42元')
    doc.add_paragraph('• 盈利：40% = 6000元')
    doc.add_paragraph('• 持仓：2天')

    add_heading(doc, '案例2：微创机器人 (02252)', 2)

    p = doc.add_paragraph()
    run = p.add_run('背景：')
    run.bold = True
    doc.add_paragraph('• 手术机器人龙头，前期涨100%')
    doc.add_paragraph('• 回调30%，机构减持')
    doc.add_paragraph('• 非恒科成分股')

    p = doc.add_paragraph()
    run = p.add_run('买入：')
    run.bold = True
    doc.add_paragraph('• 买入价：22元')
    doc.add_paragraph('• 仓位：12000元')
    doc.add_paragraph('• 触发条件：跌到整数关口20元附近企稳')

    p = doc.add_paragraph()
    run = p.add_run('催化剂：')
    run.bold = True
    doc.add_paragraph('• 医疗机器人政策利好')

    p = doc.add_paragraph()
    run = p.add_run('结果：')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 176, 80)
    doc.add_paragraph('• 卖出价：28元')
    doc.add_paragraph('• 盈利：27% = 3240元')
    doc.add_paragraph('• 持仓：5天')

    doc.add_page_break()

    # ========== 九、每日操作清单 ==========
    add_heading(doc, '九、每日操作清单', 1)

    add_heading(doc, '盘前（08:30-09:30）', 2)
    checklist1 = [
        '查看隔夜美股AI板块表现',
        '查看港股AI/机器人板块资金流向',
        '筛选跌幅榜前20（排除成分股）',
        '检查是否有前期暴涨历史',
        '查看是否有利好消息传闻',
        '确定今日候选股（最多3只）',
    ]
    for item in checklist1:
        doc.add_paragraph(f'□ {item}')

    add_heading(doc, '盘中', 2)
    checklist2 = [
        '10:00 观察候选股是否企稳',
        '10:30 确认企稳后首次建仓',
        '持续关注板块消息面',
        '有利好立即评估是否加仓',
    ]
    for item in checklist2:
        doc.add_paragraph(f'□ {item}')

    add_heading(doc, '盘后', 2)
    checklist3 = [
        '复盘今日操作',
        '更新持仓股状态',
        '检查明日是否有催化剂',
        '调整止盈止损位',
    ]
    for item in checklist3:
        doc.add_paragraph(f'□ {item}')

    # ========== 十、风险提示 ==========
    add_heading(doc, '十、风险提示', 1)

    p = doc.add_paragraph()
    run = p.add_run('⚠️ 重要提醒：')
    run.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)
    run.font.size = Pt(14)

    risks = [
        '这是高风险策略，单笔亏损可能达8%',
        '小市值股票流动性差，可能卖不出去',
        '消息面不确定，利好可能不来',
        '板块轮动快，热点可能切换',
        '建议用不超过30%资金执行此策略',
        '新手建议先模拟操作1个月',
    ]
    for i, risk in enumerate(risks, 1):
        doc.add_paragraph(f'{i}. {risk}')

    doc.add_paragraph()
    doc.add_paragraph()

    # 页脚
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f'— 堕落天使策略 · {datetime.date.today().strftime("%Y年%m月%d日")} —')
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.size = Pt(10)

    # 保存
    output_path = '/Users/mantou/hk-trading-bot/堕落天使策略_完整版.docx'
    doc.save(output_path)
    print(f'✅ 文档已保存: {output_path}')
    return output_path

if __name__ == '__main__':
    create_strategy_doc()
